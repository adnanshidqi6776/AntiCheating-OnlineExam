from flask import Flask, request, jsonify, send_from_directory
import os
import torch
import time
import csv
import cv2
from docx import Document
from datetime import datetime
from ultralytics import YOLO
import numpy as np

app = Flask(__name__)

# Load Model

device = 'cuda' if torch.cuda.is_available() else 'cpu'

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

MODEL_PATH = os.path.join(
    BASE_DIR,
    "runs",
    "detect",
    "train",
    "weights",
    "best.pt"
)

model = YOLO(MODEL_PATH)
model.to(device)

# GLOBAL STATE PER USER

user_sessions = {}

COOLDOWN = 5
CONF_THRESHOLD = 0.57

# ROUTE

@app.route("/")
def home():

    return send_from_directory("static", "home.html")

@app.route("/test")
def test_page():
    return send_from_directory("static", "testPage.html")

# CREATE SESSION

@app.route("/start_exam", methods=["POST"])
def start_exam():

    data = request.json

    nama = data["nama"]
    nim = data["nim"]
    limit = int(data["limit"])

    user_sessions[nama] = {

        "nim": nim,
        "limit": limit,
        "count": 0,
        "cooldown": False,
        "recording": False,
        "last_violation_time": 0

    }

    os.makedirs(
        f"detected_image/{nama}",
        exist_ok=True
    )

    os.makedirs(
        f"answer/{nama}",
        exist_ok=True
    )

    print("Session created:", nama)

    return jsonify({
        "status": "ok"
    })

@app.route("/detect", methods=["POST"])
def detect():

    file = request.files["image"]
    nama = request.form["nama"]

    # read image
    file_bytes = np.frombuffer(
        file.read(),
        np.uint8
    )

    frame = cv2.imdecode(
        file_bytes,
        cv2.IMREAD_COLOR
    )

    results = model(frame)

    boxes_data = []

    detected = False
    label = None
    confidence = 0

    for r in results:

        for box in r.boxes:

            cls = int(box.cls[0])

            conf = float(box.conf[0])

            name = model.names[cls]

            # FILTER CONFIDENCE
            if conf < CONF_THRESHOLD:
                continue

            x1, y1, x2, y2 = map(
                int,
                box.xyxy[0]
            )

            boxes_data.append({

                "label": name,
                "confidence": round(conf, 2),
                "bbox": [
                    x1,
                    y1,
                    x2,
                    y2
                ]

            })

            # DETECTED FLAG
            if name in [
                "handphone",
                "book",
                "finger"
            ]:

                detected = True
                label = name
                confidence = conf

    return jsonify({

        "boxes": boxes_data,
        "detected": detected,
        "label": label,
        "confidence": confidence

    })

# DETECTION EVENT

@app.route("/violation_event", methods=["POST"])
def violation_event():

    data = request.json

    nama = data["nama"]
    detected = data["detected"]
    confidence = data["confidence"]
    label = data["label"]

    session = user_sessions[nama]

    now = time.time()

    if detected:

        if session["cooldown"]:

            return jsonify({
                "action": "skip"
            })

        session["recording"] = True

        return jsonify({
            "action": "start_recording"
        })

    else:

        if session["recording"]:

            session["recording"] = False

            if now - session["last_violation_time"] >= COOLDOWN:

                session["count"] += 1

                session["cooldown"] = True

                session["last_violation_time"] = now

                log_violation(
                    nama,
                    label,
                    confidence
                )

                if session["count"] >= session["limit"]:

                    return jsonify({
                        "action": "force_end"
                    })

                return jsonify({
                    "action": "stop_recording",
                    "count": session["count"]
                })

        return jsonify({
            "action": "none"
        })


# RESET COOLDOWN

@app.route("/reset_cooldown", methods=["POST"])
def reset_cooldown():

    data = request.json

    nama = data["nama"]

    session = user_sessions[nama]

    session["cooldown"] = False

    return jsonify({
        "status": "cooldown_reset"
    })

# SAVE VIDEO EVENT

@app.route("/save_video", methods=["POST"])
def save_video():

    file = request.files["video"]
    nama = request.form["nama"]

    session = user_sessions[nama]

    count = session["count"]

    filename = f"violation_{count}.webm"

    path = f"detected_image/{nama}/{filename}"

    file.save(path)

    print("Video saved:", path)

    return jsonify({
        "status": "saved"
    })

# LOG CSV

def log_violation(nama, label, confidence):

    path = f"logs/{nama}_violation.csv"

    os.makedirs("logs", exist_ok=True)

    file_exists = os.path.isfile(path)

    with open(
        path,
        "a",
        newline=""
    ) as file:

        writer = csv.writer(file)

        if not file_exists:

            writer.writerow([
                "waktu",
                "jenis",
                "confidence"
            ])

        writer.writerow([
            datetime.now(),
            label,
            confidence
        ])

# SAVE ANSWERS

@app.route("/save_answers", methods=["POST"])
def save_answers():

    from docx import Document

    data = request.json

    nama = data["nama"]
    answers = data["answers"]

    doc = Document()

    for i, ans in enumerate(answers):

        doc.add_paragraph(
            f"Soal {i+1}"
        )

        doc.add_paragraph(ans)

        doc.add_paragraph("")

    path = f"answer/{nama}/jawaban.docx"

    doc.save(path)

    print("Answers saved")

    return jsonify({
        "status": "saved"
    })

# RUN

if __name__ == "__main__":

    app.run(
        host="0.0.0.0",
        port=5000,
        debug=True
    )